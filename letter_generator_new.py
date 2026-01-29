from docx import Document
import win32api  # still imported in case you want printing documents later
import os
import time
import pandas as pd

# -------------------------------------------------------------------------------------
# Base directory

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# -------------------------------------------------------------------------------------
# Paths

TEMPLATE_PATH = os.path.join(
    BASE_DIR,
    "apartment_letters_template",
    "template_letter.docx" # The template letter 
)

OUTPUT_DIR = os.path.join(
    BASE_DIR,
    "apartment_letters_generated"
)

RESIDENTS_FILE = os.path.join(
    BASE_DIR,
    "resident_names.xlsx"  # Excel with Apartment and Resident Name
)

# -------------------------------------------------------------------------------------
# Create output folder if it doesn't exist

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# -------------------------------------------------------------------------------------
# Load resident names

df_residents = pd.read_excel(RESIDENTS_FILE)

# Replace NaN (empty Excel cells) with empty strings
df_residents["Resident Name"] = df_residents["Resident Name"].fillna("")

# Create dictionary: { "A1": "John Doe", ... }
resident_dict = dict(
    zip(df_residents["Apartment"], df_residents["Resident Name"])
)

# -------------------------------------------------------------------------------------
# Apartment details

blocks = ["A", "B", "C"]
units_per_block = 10

# -------------------------------------------------------------------------------------
# Function to replace text in paragraphs and tables

def replace_text(doc, old_text, new_text):
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)

# -------------------------------------------------------------------------------------
# Generate letters

for block in blocks:
    for unit in range(1, units_per_block + 1):
        apartment = f"{block}{unit}"

        # Get resident name (may be empty)
        resident_name = resident_dict.get(apartment, "").strip()

        # Build greeting
        if resident_name:
            greeting = f"Dear {resident_name},"
        else:
            greeting = f"Dear resident of apartment {apartment},"

        # Load the template
        doc = Document(TEMPLATE_PATH)

        # Replace greeting in template
        replace_text(
            doc,
            "Dear *name*,",
            greeting
        )

        # Output file path
        output_file = os.path.join(
            OUTPUT_DIR,
            f"Letter_Apartment_{apartment}.docx"
        )

        # Save document
        doc.save(output_file)

        print(f"Generated letter for apartment {apartment}")

        time.sleep(1)  # avoids IO overload

# -------------------------------------------------------------------------------------
# Done
print("All letters generated successfully.")

